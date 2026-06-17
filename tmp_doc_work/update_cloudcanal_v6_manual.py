from pathlib import Path
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement

SRC = Path('/Users/pudding/Documents/CloudCanalV6使用说明书.docx')
OUT = Path('/Users/pudding/Documents/CloudCanalV6使用说明书_软著更新.docx')


def set_para_text(paragraph, text):
    p = paragraph._p
    pPr = p.pPr
    for child in list(p):
        if child is not pPr:
            p.remove(child)
    paragraph.add_run(text)


def insert_paragraph_after(doc, idx, text, style_name=None):
    src_p = doc.paragraphs[idx]
    new_xml = deepcopy(src_p._p)
    src_p._p.addnext(new_xml)
    new_para = doc.paragraphs[idx + 1]
    for r in list(new_para.runs):
        r.text = ''
    if style_name:
        new_para.style = style_name
    set_para_text(new_para, text)
    return new_para


def norm(text):
    return ' '.join((text or '').replace('\xa0', ' ').split())

def find_para(doc, exact=None, contains=None, start=0):
    exact_n = norm(exact) if exact is not None else None
    contains_n = norm(contains) if contains is not None else None
    for i, p in enumerate(doc.paragraphs[start:], start=start):
        t = norm(p.text.strip())
        if exact_n is not None and t == exact_n:
            return i
        if contains_n is not None and contains_n in t:
            return i
    return -1


missing = []

def replace_contains(doc, contains, text):
    i = find_para(doc, contains=contains)
    if i < 0:
        missing.append(contains)
        return -1
    set_para_text(doc.paragraphs[i], text)
    return i

def replace_all_contains(doc, contains, text):
    hits = []
    start = 0
    while True:
        i = find_para(doc, contains=contains, start=start)
        if i < 0:
            break
        hits.append(i)
        set_para_text(doc.paragraphs[i], text)
        start = i + 1
    if not hits:
        missing.append(contains)
    return hits


def main():
    doc = Document(SRC)
    set_para_text(doc.paragraphs[0], '开云集致CloudCanal数据融合软件[简称：CloudCanal]V6.0')
    set_para_text(doc.paragraphs[1], '使用说明书')
    if len(doc.paragraphs) > 2 and doc.paragraphs[2].text.strip() == '目类':
        set_para_text(doc.paragraphs[2], '目录')

    replacements = [
        ('CloudCanal 是一款 数据同步、迁移 工具', 'CloudCanal 是一款面向企业级数据融合场景的数据迁移、实时同步、数据校验、数据订正与数据加工工具，帮助企业构建高质量数据管道，具备实时高效、精确互联、稳定可扩展、一站式、混合部署、复杂数据转换、可视化运维和开放集成等特点。'),
        ('可选搭配结构迁移、数据校验和订正，满足结构准备或数据质量的需求。', '可选搭配结构迁移、数据校验、数据订正、目标表重建、全量前清空目标数据和周期调度等能力，满足结构准备、初始化迁移、数据质量复核和持续运维需求。'),
        ('可选搭配结构迁移、数据初始化(全量迁移)、单次或定时数据校验与订正', '可选搭配结构迁移、数据初始化（全量迁移）、DDL 同步、修改订阅、单次或定时数据校验与订正，满足数据准备和业务长周期数据同步对于数据质量、结构变更和链路稳定性的要求。'),
        ('将源端和对端数据分别取出，逐字段对比', '将源端和目标端数据分别取出，按主键、唯一键或指定字段逐行逐列对比，可选择差异数据订正。该功能可单独使用，也可配合结构迁移、全量迁移、增量同步使用，满足用户数据质量验证与修复需求。'),
        ('CloudCanal 目前提供四种版本的产品', 'CloudCanal 目前提供社区版、SaaS 版、商业试用版和商业版四种产品形态，支持私有部署、全托管和 BYOC 等使用方式，用户可根据任务规模、部署环境、安全要求和服务需求选择相应版本。'),
        ('SaaS 版：BYOC 部署，功能与商业版一致，轻量化部署', 'SaaS 版：支持全托管和 BYOC 方式，功能与商业版一致，可降低控制台运维成本，适用于个人项目、中小型团队以及需要快速构建数据链路的生产场景。费用参考产品定价。'),
        ('在支持的数据源方面，商业版在社区版基础上，还支持', '在支持的数据源方面，CloudCanal V6 已覆盖 60+ 数据源类型，包含 MySQL、PostgreSQL、Oracle、SQL Server、MongoDB、Redis、Kafka、RocketMQ、RabbitMQ、Doris、StarRocks、ClickHouse、Hologres、Snowflake、OceanBase、KingbaseES、openGauss、达梦、DB2、TiDB、MariaDB、DynamoDB 等数据库、消息系统、缓存、湖仓和云托管数据源，商业版本按授权范围提供更完整的数据源链路能力。'),
        ('CloudCanal 商业版支持任务组功能。当前任务组分为业务组和并行组两种。', 'CloudCanal 支持任务组能力，当前任务组分为业务组和并行组两种。业务组可帮助用户查看并管理具有业务关联性的一组任务；并行组可使用多个任务对相同的表数据进行分区迁移同步，当迁移同步海量数据时可有效提升同步效率。'),
        ('支持 MySQL、PostgreSQL、Oracle、Kafka、SQL Server、MongoDB、Redis、SAP HANA、TiDB 等多种同构数据源之间的数据迁移同步', '支持 MySQL、PostgreSQL、Oracle、Kafka、SQL Server、MongoDB、Redis、SAP HANA、TiDB、OceanBase、KingbaseES、openGauss、达梦、Doris、StarRocks、ClickHouse 等多种同构或异构数据源之间的数据迁移同步，为数据库版本升级、多云/云上云下数据同步、跨区域数据迁移同步以及容灾数据同步带来便利。'),
        ('汇集结构迁移、数据迁移、数据同步、数据校验与订正、修改订阅等功能', '汇集结构迁移、全量迁移、增量同步、DDL 同步、数据校验与订正、修改订阅、任务告警、操作审计和开放 API 等功能，通过有限状态机让功能自动流转和运行。一站式支持用户数据准备、长期同步、质量复核与运维治理过程。'),
        ('参考 Linux/Mac 安装文档，在官网下载 CloudCanal', '参考 Docker、Kubernetes 或 TGZ 安装文档，在官网下载或获取 CloudCanal 安装包，并按部署环境准备 JDK17、元数据库、网络访问和运行目录。'),
        ('安装完整产品(Linux/MacOS)。', '安装完整产品，可选择 Docker、Kubernetes 或 TGZ 部署方式；TGZ 部署需确认 JDK17 环境，Docker/Kubernetes 部署按镜像和编排配置启动控制台、Sidecar 与 Worker 相关服务。'),
        ('选择任务类型为 增量同步 并勾选 全量初始化。', '选择任务类型为 增量同步，并按需要勾选 全量初始化、结构迁移、DDL 同步、数据校验、数据订正或数据加工等功能。'),
        ('选择任务类型为 增量同步，并勾选 全量初始化, 点击 下一步。', '选择任务类型为 增量同步，并按业务需求勾选 全量初始化、结构迁移、DDL 同步、数据校验、数据订正或数据加工，点击 下一步。'),
        ('CloudCanal 支持创建全量迁移和增量同步一体化任务', 'CloudCanal 支持创建结构迁移、全量迁移和增量同步一体化任务，包含表结构迁移、全量数据初始化、增量实时同步、DDL 同步、数据校验、数据订正等多个阶段，整个过程可在创建向导中配置并由系统自动流转。增量实时同步阶段，任务通过启动位点控制确保全量期间产生的增量也完整同步到目标端。'),
        ('任务详情页支持在线查看日志，让运维任务更加简单、快捷。', '任务详情页支持在线查看结构迁移、全量迁移、增量同步、数据校验、数据订正等阶段日志，让运维人员快速确认任务实时运行状态、写入结果、异常堆栈和问题定位线索。'),
        ('CloudCanal 支持用户修改任务的运行参数。任务参数与任务运行的性能、工作机制息息相关。', 'CloudCanal 支持用户修改任务运行参数。任务参数与任务运行性能、错误处理、并发能力、写入策略、日志输出、位点管理和任务稳定性息息相关。'),
        ('CloudCanal 控制台可以直接查看机器的监控指标，并且在机器状态异常时发送报警。', 'CloudCanal 控制台可以直接查看机器和任务监控指标，并且在机器状态异常、任务中断、延迟升高、连接失败、资源异常或数据不一致时发送告警。用户排查任务卡顿问题时可结合任务指标、机器监控、告警日志和异常日志共同分析。'),
        ('CloudCanal 提供告警日志，以便用户清晰地掌握告警历史。', 'CloudCanal 提供告警日志，以便用户清晰掌握告警历史、发送结果、异常类型和处理状态。V6.1 强化任务告警能力，支持批量配置、告警频率控制、异常告警白名单过滤以及单任务绑定多个 IM webhook 告警地址。'),
        ('CloudCanal 支持 自建数据源 和 云托管数据源。', 'CloudCanal 支持自建数据源、云托管数据源、消息系统、缓存、湖仓、文件和云服务等多种数据源类型。用户可根据部署位置选择内网、外网、SSH 隧道、TLS/SSL、云厂商 AK/SK 或 BYOC 网络方式接入。'),
        ('目前支持的数据库包括', '目前支持的数据源包括 AutoMQ、ClickHouse、Dameng（达梦）、Db2、Doris、DynamoDB、Elasticsearch、GaussDB、Greenplum、Hana、Hive、Hudi、Kafka、KingbaseES、Kudu、MariaDB、MongoDB、MySQL、OceanBase、OpenGauss、Oracle、PolarDB、PostgreSQL、Pulsar、RabbitMQ、Redis、RocketMQ、SelectDB、Snowflake、SQL Server、StarRocks、TiDB、Tunnel、VastBase 等，具体支持版本以数据源种类和版本支持文档为准。'),
        ('当前同时支持 主子账号 单独设置，认证应用支持 Microsoft Authenticator 或 Google Authenticator 。', '当前同时支持主子账号单独设置，认证应用支持 Microsoft Authenticator 或 Google Authenticator。管理员还可结合角色、资源权限、SSO、LDAP、OIDC、钉钉、飞书、企业微信等统一身份认证能力提升访问安全。'),
        ('CloudCanal 权限由两部分组成，即功能权限和资源权限。', 'CloudCanal 权限由功能权限和资源权限两部分组成。功能权限控制用户可访问的菜单和操作，资源权限控制用户可访问的数据源、同步任务等资源范围，便于企业按团队、项目、环境和岗位分配权限。'),
        ('CloudCanal 支持对数据源设置自定义的环境，以区分数据源的业务场景', 'CloudCanal 支持对数据源设置自定义环境，以区分测试、预发、生产、灾备、云上、云下等业务场景，方便用户在创建数据源、创建任务、搜索资源和审计操作时进行识别和管理。'),
    ]
    for old, new in replacements:
        replace_contains(doc, old, new)

    replace_all_contains(doc, '选择任务类型为 增量同步，并勾选 全量初始化, 点击 下一步。', '选择任务类型为 增量同步，并按业务需求勾选 全量初始化、结构迁移、DDL 同步、数据校验、数据订正或数据加工，点击 下一步。')

    anchor = find_para(doc, exact='纵向为源端，横向为目标端（因数据源较多，分为两部分）')
    if anchor >= 0 and find_para(doc, exact='1.1.1 V6版本能力说明') < 0:
        insert_paragraph_after(doc, anchor, '1.1.1 V6版本能力说明', 'Heading 3')
        insert_paragraph_after(doc, anchor + 1, 'CloudCanal V6.0 版本产品运行环境升级至 JDK17，继续增强结构迁移、全量迁移、增量同步、数据校验、数据订正、DDL 同步和数据加工能力，并新增或增强 MySQL 到 Hologres、PostgreSQL 到 CloudBerry、GreenPlum 到 CloudBerry、MySQL/PostgreSQL/Oracle 到 Snowflake、KingbaseES 与 OceanBase for Oracle 互通、OpenGauss 到 Doris、MySQL/PolarDB MySQL 到 DuckDB 等数据链路。', 'Normal')
        insert_paragraph_after(doc, anchor + 2, 'CloudCanal V6.1 在 V6.0 基础上强化任务告警、批量告警配置、告警频率和异常告警白名单过滤，支持 Kafka 源端扁平化消息格式消费、OceanBase MySQL Binlog GTID 位点、DynamoDB 多 Schema 迁移同步、TiDB SSL 连接、单任务绑定多个 IM webhook 告警地址，并优化超多表任务创建和修改订阅时的搜索性能。', 'Normal')

    alarm_heading = find_para(doc, exact='6.4查看告警日志')
    if alarm_heading >= 0 and find_para(doc, contains='V6.1 强化任务告警能力') < 0:
        insert_paragraph_after(doc, alarm_heading, 'CloudCanal 提供告警日志，以便用户清晰掌握告警历史、发送结果、异常类型和处理状态。V6.1 强化任务告警能力，支持批量配置、告警频率控制、异常告警白名单过滤以及单任务绑定多个 IM webhook 告警地址。', 'Normal')

    before_protocol = find_para(doc, exact='协议')
    if before_protocol > 0 and find_para(doc, exact='6.9开放 API 与平台集成') < 0:
        insert_paragraph_after(doc, before_protocol - 1, '6.9开放 API 与平台集成', 'Heading 2')
        insert_paragraph_after(doc, before_protocol, 'CloudCanal 提供基于 HTTP 的开放 API，支持使用 AK/SK 认证和 RequestId 调用链跟踪，外部系统可通过接口查询和管理数据源、任务、Worker、集群、用户、常量和异步任务等对象。', 'Normal')
        insert_paragraph_after(doc, before_protocol + 1, '调用方在集成前应先获取 AK/SK，查询可用集群和数据源，测试所选集群到源端、目标端数据源的连通性，再构建库、表、列元数据和映射关系，最后创建、启动、停止、查询或删除数据任务。API 凭证应妥善保存，并按最小权限原则分配给业务平台或调度系统。', 'Normal')

    doc.save(OUT)
    print(OUT)
    if missing:
        print('MISSING_MATCHES=' + str(len(missing)))
        for item in missing:
            print('- ' + item)

if __name__ == '__main__':
    main()
